"""

================================================================================

  IRCTC Schedule Generator - Suhail Edition v2.0 (Unlimited)

  Backend: FastAPI

  Source : IRCTC Official JSON API

  Output : IRCTC_Tender_Schedules.docx

  IRCTC Flow per train:

    GET /trnscheduleenquiry/{train_no} -> JSON Schedule

================================================================================

"""

import os

import sys

import re

import asyncio

import io

try:

    if hasattr(sys.stdout, 'reconfigure'):

        sys.stdout.reconfigure(encoding='utf-8', errors='ignore')

    if hasattr(sys.stderr, 'reconfigure'):

        sys.stderr.reconfigure(encoding='utf-8', errors='ignore')

except Exception:

    pass

import json

import uuid

import time

import tempfile

from typing import List, Optional, Tuple

import requests

import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks

from fastapi.middleware.cors import CORSMiddleware

from fastapi.responses import FileResponse, StreamingResponse

from fastapi.staticfiles import StaticFiles

from pydantic import BaseModel

import pdfplumber

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

from docx import Document

from docx.shared import Cm, Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

# Global lazy OCR reader (loads ~300MB model on first use, then cached in RAM)

_ocr_reader = None

def _get_ocr_reader():

    """Lazily initialise easyocr.Reader once and reuse across requests."""

    global _ocr_reader

    if _ocr_reader is None:

        print("[OCR] Loading easyocr model (first-time only, please wait)...")

        import easyocr

        _ocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)

        print("[OCR] Model loaded.")

    return _ocr_reader

# ==============================================================================

# SECTION 1 - CONFIG & APP SETUP

# ==============================================================================

_DAY_NAMES = ["MON", "TUE", "WED", "THU", "FRI", "SAT", "SUN"]

app = FastAPI(title="IRCTC Schedule Generator - Suhail Edition v2.1")

app.add_middleware(

    CORSMiddleware,

    allow_origins=["*", "null"],

    allow_credentials=True,

    allow_methods=["*"],

    allow_headers=["*"],

)

# ==============================================================================

# SECTION 2 - INDIARAILINFO LIVE SCRAPER ENGINE

# ==============================================================================

import urllib.parse

from bs4 import BeautifulSoup

from curl_cffi import requests as curl_requests

os.system('')  # Enable ANSI terminal colors in Windows CMD / PowerShell

class C:

    RESET      = "\033[0m"

    BOLD       = "\033[1m"

    DIM        = "\033[2m"

    RED        = "\033[31m"

    GREEN      = "\033[32m"

    YELLOW     = "\033[33m"

    BLUE       = "\033[34m"

    MAGENTA    = "\033[35m"

    CYAN       = "\033[36m"

    WHITE      = "\033[37m"

    

    B_RED      = "\033[91m"

    B_GREEN    = "\033[92m"

    B_YELLOW   = "\033[93m"

    B_BLUE     = "\033[94m"

    B_MAGENTA  = "\033[95m"

    B_CYAN     = "\033[96m"

    B_WHITE    = "\033[97m"

print(f"\n{C.B_CYAN}{C.BOLD}=" * 65 + f"{C.RESET}")

print(f"  {C.B_YELLOW}{C.BOLD}🚂 IRCTC TENDER SCHEDULE GENERATOR ENGINE  v3.0{C.RESET}  {C.B_GREEN}● ONLINE{C.RESET}")

print(f"  {C.B_CYAN}⚡ IndiaRailInfo Scraper Initialised | No Cookies/Tokens Required{C.RESET}")

print(f"  {C.B_MAGENTA}🎨 Vibrant ANSI Console Output Active{C.RESET}")

print(f"{C.B_CYAN}{C.BOLD}=" * 65 + f"{C.RESET}\n")

_CLUTTER_PARENS_RE = re.compile(

    r"\s*\((?:PT\d*|SF|Mail|Express|Special|UnReserved|Weekly|Bi-Weekly|Tri-Weekly|Daily|AC|TOD|TOD\+WCB|WCB)\)",

    re.IGNORECASE

)

def _clean_train_name(name: str) -> str:

    if not name:

        return ""

    name = _CLUTTER_PARENS_RE.sub("", name)

    name = re.sub(r"\s+", " ", name).strip()

    return name

def _extract_id_from_url(url: str, train_no: str) -> Optional[str]:

    if not url:

        return None

    m = re.search(r"indiarailinfo[.]com/train/(.+)", url)

    if not m:

        return None

    path = m.group(1).split("?")[0]

    segments = [s for s in path.split("/") if s]

    nums = [s for s in segments if s.isdigit()]

    if not nums:

        return None

    if len(nums) == 1:

        return nums[0]

    for num in nums:

        if num != train_no and len(num) >= 3:

            return num

    return nums[0]

# ==============================================================================
# SECTION 1.5 - OFFLINE MASTER DATASET FAIL-SAFE CACHE (3,596+ TRAINS)
# ==============================================================================

MASTER_DB_CACHE: dict = {}

def _load_master_db():
    global MASTER_DB_CACHE
    possible_paths = [
        os.path.join(os.path.dirname(__file__), "all_india_train_pairs_master.json"),
        os.path.join(os.path.dirname(os.path.dirname(__file__)), "all_india_train_pairs_master.json"),
        os.path.join(os.getcwd(), "all_india_train_pairs_master.json"),
        os.path.join(os.getcwd(), "backend", "all_india_train_pairs_master.json")
    ]
    for p in possible_paths:
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    for pair in data.get("train_pairs", []):
                        if pair.get("up_train"):
                            up = pair["up_train"]
                            MASTER_DB_CACHE[str(up["number"]).strip()] = {
                                "train_number": str(up["number"]).strip(),
                                "train_name": up.get("name", f"TRAIN {up['number']}"),
                                "origin_code": up.get("origin_code", ""),
                                "origin_name": up.get("origin_name", ""),
                                "dest_code": up.get("dest_code", ""),
                                "dest_name": up.get("dest_name", ""),
                                "running_days": up.get("running_days", "(DAILY)"),
                                "slug": up.get("slug", str(up["number"])),
                                "train_type": pair.get("train_type", "Express"),
                                "zone": pair.get("zone", "IR")
                            }
                        if pair.get("down_train"):
                            dn = pair["down_train"]
                            MASTER_DB_CACHE[str(dn["number"]).strip()] = {
                                "train_number": str(dn["number"]).strip(),
                                "train_name": dn.get("name", f"TRAIN {dn['number']}"),
                                "origin_code": dn.get("origin_code", ""),
                                "origin_name": dn.get("origin_name", ""),
                                "dest_code": dn.get("dest_code", ""),
                                "dest_name": dn.get("dest_name", ""),
                                "running_days": dn.get("running_days", "(DAILY)"),
                                "slug": dn.get("slug", str(dn["number"])),
                                "train_type": pair.get("train_type", "Express"),
                                "zone": pair.get("zone", "IR")
                            }
                print(f"[RENDER LOG] 🛡️ Master Train Database Loaded: {len(MASTER_DB_CACHE)} trains in offline fail-safe cache.", flush=True)
                return
            except Exception as e:
                print(f"[RENDER LOG] ⚠️ Error loading master database from {p}: {e}", flush=True)

_load_master_db()

IRI_CDN_NODES = [
    "srv1.indiarailinfo.com",
    "srv3.indiarailinfo.com",
    "srv2.indiarailinfo.com",
    "m.indiarailinfo.com",
    "indiarailinfo.com"
]

def _solve_challenge_on_session(session: curl_requests.Session, html: str, domain: str) -> bool:
    """Solve IndiaRailInfo browser verification challenge automatically."""
    try:
        sig_match = re.search(r"data-sig=['\"]([^'\"]+)['\"]", html)
        if sig_match:
            sig = sig_match.group(1)
            x_val = sig.split("|")[-1] if "|" in sig else "55"
            token = f"0:5:2:1:8:1:1:0:{x_val}:{sig}:0"
            verify_url = f"https://{domain}/verify-browser?t={token}"
            rv = session.get(verify_url, timeout=8)
            print(f"[RENDER LOG] 🔐 Solved Challenge on {domain} -> HTTP {rv.status_code} ({rv.text.strip()[:30]})", flush=True)
            time.sleep(0.4)
            return rv.status_code == 200
    except Exception as e:
        print(f"[RENDER LOG] ⚠️ Challenge solver exception on {domain}: {e}", flush=True)
    return False

def _get_active_iri_session() -> tuple[curl_requests.Session, str]:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9,hi;q=0.8",
        "Sec-Ch-Ua": '"Chromium";v="124", "Google Chrome";v="124", "Not-A.Brand";v="99"',
        "Sec-Ch-Ua-Mobile": "?0",
        "Sec-Ch-Ua-Platform": '"Windows"',
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "same-origin",
        "Upgrade-Insecure-Requests": "1",
    }
    for node in IRI_CDN_NODES:
        try:
            session = curl_requests.Session(impersonate="chrome124")
            session.headers.update(headers)
            session.headers["Referer"] = f"https://{node}/"
            r = session.get(f"https://{node}/", timeout=6)
            print(f"[RENDER LOG] 🌐 Testing CDN Node: {node:25s} -> HTTP {r.status_code} (len: {len(r.text)})", flush=True)
            if r.status_code == 200:
                if "data-sig" in r.text or "iri-xsig" in r.text:
                    _solve_challenge_on_session(session, r.text, node)
                print(f"[RENDER LOG] ✅ Connected & Verified on Node: https://{node}/", flush=True)
                return session, node
        except Exception as e:
            print(f"[RENDER LOG] ⚠️ Node {node} check error: {e}", flush=True)
            continue
    print("[RENDER LOG] ⚠️ Fallback to default node: srv1.indiarailinfo.com", flush=True)
    def_session = curl_requests.Session(impersonate="chrome124")
    def_session.headers.update(headers)
    return def_session, "srv1.indiarailinfo.com"

def _solve_iri_challenge_and_fetch(url: str, session: curl_requests.Session) -> str:
    domain = urllib.parse.urlparse(url).netloc
    try:
        r = session.get(url, timeout=15)
        if ("data-sig" in r.text or "iri-xsig" in r.text) and len(r.text) < 5000:
            print(f"[RENDER LOG] 🔄 Solving inline challenge for {url}...", flush=True)
            _solve_challenge_on_session(session, r.text, domain)
            r = session.get(url, timeout=15)
        return r.text
    except Exception as e:
        print(f"[RENDER LOG] ❌ IRI fetch exception for {url}: {e}", flush=True)
        return ""

def _extract_id_from_url(url: str, train_no: str) -> Optional[str]:
    if not url:
        return None
    patterns = [
        r'/train/[^/]+/(\d+)',
        r'/train/(\d+)',
        r'indiarailinfo\.com/train/[^/]+/(\d+)',
        r'indiarailinfo\.com/train/(\d+)'
    ]
    for pat in patterns:
        m = re.search(pat, url)
        if m:
            candidate = m.group(1)
            if candidate != train_no and len(candidate) >= 3:
                return candidate
    nums = re.findall(r'\b\d{3,7}\b', url)
    for num in nums:
        if num != train_no and len(num) >= 3:
            return num
    return None

def _resolve_iri_slug(train_no: str, session: curl_requests.Session, node: str = "srv1.indiarailinfo.com") -> str:
    train_no = str(train_no).strip()
    
    # Strategy 0: Instant Master Dataset Slug Shortcut
    if train_no in MASTER_DB_CACHE and MASTER_DB_CACHE[train_no].get("slug"):
        master_slug = MASTER_DB_CACHE[train_no]["slug"]
        if master_slug and master_slug.isdigit() and len(master_slug) >= 3:
            print(f"[RENDER LOG] ⚡ Instant Master DB Slug for {train_no} -> {master_slug}", flush=True)
            return master_slug
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
        "Referer": f"https://{node}/"
    }

    # Strategy 1: IndiaRailInfo Official Direct Internal API
    try:
        url_internal = f"https://{node}/shtml/list.shtml?LappGetTrainList/{train_no}/0/0/0"
        r_int = session.get(url_internal, headers=headers, timeout=6)
        print(f"[RENDER LOG] 🔍 Internal API query: {url_internal} -> HTTP {r_int.status_code} (len: {len(r_int.text)})", flush=True)
        
        if ("data-sig" in r_int.text or "iri-xsig" in r_int.text) and not ("dropdowntable" in r_int.text):
            print(f"[RENDER LOG] 🔐 Solving challenge on list.shtml for {node}...", flush=True)
            _solve_challenge_on_session(session, r_int.text, node)
            r_int = session.get(url_internal, headers=headers, timeout=6)

        if r_int.status_code == 200 and "dropdowntable" in r_int.text:
            soup_int = BeautifulSoup(r_int.text, "html.parser")
            for tr in soup_int.find_all("tr", class_=re.compile(r"rowM1", re.IGNORECASE)):
                tds = tr.find_all("td")
                if tds and tds[0].get_text().strip().isdigit():
                    slug_id = tds[0].get_text().strip()
                    print(f"[RENDER LOG] 🎯 Direct API Resolved Train {train_no} -> Slug ID: {slug_id}", flush=True)
                    return slug_id
    except Exception as e:
        print(f"[RENDER LOG] ⚠️ Internal API slug strategy error for {train_no}: {e}", flush=True)

    # Strategy 2: DuckDuckGo HTML Search Indexing
    try:
        url_ddg = f"https://html.duckduckgo.com/html/?q=site:indiarailinfo.com+train+{train_no}"
        r_ddg = session.get(url_ddg, headers=headers, timeout=6)
        if r_ddg.status_code == 200:
            soup_ddg = BeautifulSoup(r_ddg.text, "html.parser")
            for a in soup_ddg.find_all("a"):
                href = a.get("href", "")
                if "indiarailinfo.com/train/" in href or "uddg=" in href:
                    unquoted = urllib.parse.unquote(href)
                    tid = _extract_id_from_url(unquoted, train_no)
                    if tid:
                        print(f"[RENDER LOG] 🎯 DuckDuckGo Resolved Train {train_no} -> Slug ID: {tid}", flush=True)
                        return tid
    except Exception as e:
        print(f"[RENDER LOG] ⚠️ DDG search slug strategy error for {train_no}: {e}", flush=True)

    # Strategy 3: Yahoo Search Indexing
    try:
        url_y = f"https://search.yahoo.com/search?p=site:indiarailinfo.com+train+{train_no}"
        r_y = session.get(url_y, headers=headers, timeout=6)
        if r_y.status_code == 200:
            soup_y = BeautifulSoup(r_y.text, "html.parser")
            for a in soup_y.find_all("a"):
                href = a.get("href", "")
                if "RU=" in href:
                    unquoted = urllib.parse.unquote(href)
                    tid = _extract_id_from_url(unquoted, train_no)
                    if tid:
                        print(f"[RENDER LOG] 🎯 Yahoo Resolved Train {train_no} -> Slug ID: {tid}", flush=True)
                        return tid
    except Exception as e:
        print(f"[RENDER LOG] ⚠️ Yahoo search slug strategy error for {train_no}: {e}", flush=True)

    print(f"[RENDER LOG] ℹ️ Using default train_no as slug: {train_no}", flush=True)
    return train_no

def fetch_train_indiarailinfo(train_no: str) -> Optional[dict]:
    train_no = str(train_no).strip()
    print(f"  {C.B_CYAN}🔍 [IRI]{C.RESET} Fetching Schedule for Train {C.B_YELLOW}{train_no}{C.RESET}...", flush=True)

    session, node = _get_active_iri_session()
    train_id = _resolve_iri_slug(train_no, session, node)
    iri_url = f"https://{node}/train/{train_id}"
    print(f"  {C.CYAN}🔗 [IRI Source URL]{C.RESET} {C.DIM}-> {iri_url}{C.RESET}", flush=True)

    try:
        html = _solve_iri_challenge_and_fetch(iri_url, session)
        if not html or len(html) < 5000:
            print(f"[RENDER LOG] ⚠️ HTML size ({len(html)} bytes) too small. Retrying in 1s...", flush=True)
            time.sleep(1)
            html = _solve_iri_challenge_and_fetch(iri_url, session)

        print(f"[RENDER LOG] 📥 HTML fetched: {len(html)} bytes for Train {train_no}", flush=True)
        if not html or len(html) < 5000:
            print(f"[RENDER LOG] ⚠️ Online scraping failed for Train {train_no}. Checking Offline Master Cache...", flush=True)
            if train_no in MASTER_DB_CACHE:
                cached = MASTER_DB_CACHE[train_no]
                print(f"[RENDER LOG] 🛡️ [OFFLINE FAIL-SAFE ACTIVATED] Loaded Train {train_no} from Master Database: {cached['train_name']} ({cached['origin_code']} -> {cached['dest_code']})", flush=True)
                return {
                    "train_number":  train_no,
                    "train_name":    cached["train_name"],
                    "origin_code":   cached["origin_code"] or "SRC",
                    "dest_code":     cached["dest_code"] or "DST",
                    "dep_time":      "---",
                    "arr_time":      "---",
                    "station_codes": [cached["origin_code"], cached["dest_code"]] if cached["origin_code"] and cached["dest_code"] else ["SRC", "DST"],
                    "coaches":       "20 Coaches",
                    "run_days":      cached["running_days"]
                }
            print(f"  {C.B_RED}⚠️ [WARN]{C.RESET} Train {C.B_YELLOW}{train_no}{C.RESET} not found on IndiaRailInfo or Offline Master Cache.", flush=True)
            return None

        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string if soup.title else f"TRAIN {train_no}"

        # Train Name
        train_name = f"TRAIN {train_no}"
        if "/" in title:
            after_slash = title.split("/", 1)[1].strip()
            parts = after_slash.split(" - ")
            name_parts = []
            for p in parts:
                if re.search(r"\b\w+\s+to\s+\w+\b", p, re.IGNORECASE) or "Railway Enquiry" in p or "Zone" in p:
                    break
                name_parts.append(p.strip())
            if name_parts:
                train_name = _clean_train_name(" - ".join(name_parts))

        # Station Codes in exact order
        station_codes = []
        for a in soup.find_all("a", href=re.compile(r"/station/map/.*#st")):
            txt = a.get_text().strip()
            if 2 <= len(txt) <= 5 and txt.isupper() and txt not in station_codes:
                station_codes.append(txt)

        if len(station_codes) < 2:
            for a in soup.find_all("a", href=re.compile(r"/station/")):
                txt = a.get_text().strip()
                if "/" in txt:
                    code = txt.split("/")[0].strip()
                    if 2 <= len(code) <= 5 and code.isupper() and code not in station_codes:
                        station_codes.append(code)

        if len(station_codes) < 2:
            print(f"[RENDER LOG] ❌ Insufficient station codes ({len(station_codes)}) found for train {train_no}", flush=True)
            return None

        origin_code = station_codes[0]
        dest_code = station_codes[-1]

        # Safeguard against identical origin and destination codes
        if origin_code == dest_code and len(station_codes) > 1:
            for c in reversed(station_codes):
                if c != origin_code:
                    dest_code = c
                    break

        # Departure & Arrival Times
        dep_m = re.search(r"Departs\s*@\s*(?:<[^>]+>\s*)*([0-2]?\d:[0-5]\d)", html, re.IGNORECASE | re.DOTALL)
        arr_m = re.search(r"Arrives\s*@\s*(?:<[^>]+>\s*)*([0-2]?\d:[0-5]\d)", html, re.IGNORECASE | re.DOTALL)

        dep_time = "---"
        if dep_m:
            parts = dep_m.group(1).split(":")
            dep_time = f"{int(parts[0]):02d}{int(parts[1]):02d} hrs"

        arr_time = "---"
        if arr_m:
            parts = arr_m.group(1).split(":")
            arr_time = f"{int(parts[0]):02d}{int(parts[1]):02d} hrs"

        # Running Days Bitmask
        run_days_str = "(DAILY)"
        dep_tag = soup.find(string=re.compile(r"Departs\s*@", re.IGNORECASE))
        if dep_tag and dep_tag.parent:
            dep_cell = dep_tag.parent
            grid = dep_cell.find("table", class_=re.compile(r"deparrgrid"))
            if grid:
                tds = grid.find_all("td")
                if len(tds) == 7:
                    bits = []
                    for td in tds:
                        style = str(td.get("style", ""))
                        if "opacity" in style and ("0.2" in style or "0.3" in style or "0.4" in style):
                            bits.append("0")
                        else:
                            bits.append("1")
                    mon_sun_bits = bits[1:] + bits[:1]
                    run_days_str = _format_run_days("".join(mon_sun_bits))

        # Coaches
        coaches_str = "20 Coaches"
        coach_match = re.search(r"(\d+)\s*(?:LHB|ICF)?\s*Coaches", html, re.IGNORECASE)
        if coach_match:
            coaches_str = f"{coach_match.group(1)} Coaches"

        print(f"[RENDER LOG] 🚉 Train {train_no} Parsed: {train_name} | {origin_code} -> {dest_code} | Dep: {dep_time} | Arr: {arr_time} | Stations: {len(station_codes)}", flush=True)
        print(f"  {C.B_GREEN}✅ [OK]{C.RESET} Successfully fetched Train {C.B_YELLOW}{train_no}{C.RESET} ({C.B_WHITE}{train_name}{C.RESET}) [{dep_time} - {arr_time}]", flush=True)

        return {
            "train_number":  train_no,
            "train_name":    train_name,
            "origin_code":   origin_code,
            "dest_code":     dest_code,
            "dep_time":      dep_time,
            "arr_time":      arr_time,
            "run_days":      run_days_str,
            "station_codes": station_codes,
            "coaches":       coaches_str,
        }

    except Exception as e:
        print(f"[RENDER LOG] ❌ IndiaRailInfo scraper error for train {train_no}: {e}", flush=True)
        return None

def _suffix_clean(name: str) -> str:

    return _SUFFIX_RE.sub("", name).strip()

def _normalise_time(raw: str) -> str:

    """

    Convert times to 'HHMM hrs'.

    Supports dot-separated '20.10' or colon-separated '20:10'.

    """

    if not raw:

        return "---"

    raw = raw.strip()

    if raw.lower() in ("first", "last", "source", "destination",

                       "--", "---", "--:--", "--.--", "0", ""):

        return "---"

    normalised = raw.replace(".", ":")

    parts = normalised.split(":")

    if len(parts) >= 2:

        try:

            hh = int(parts[0])

            mm = int(parts[1])

            return f"{hh:02d}{mm:02d} hrs"

        except ValueError:

            pass

    return raw

def _format_run_days(bitmask: str) -> str:

    """

    Conditional frequency formatting:

      7 days  -> (DAILY)

      4 or 5  -> 04 DAYS (Except – MON, THU, SUN)   [show missing days]

      1,2,3,6 -> 03 DAYS (TUE, THU, SUN)             [show running days]

    Index: 0=MON 1=TUE 2=WED 3=THU 4=FRI 5=SAT 6=SUN

    """

    active  = [_DAY_NAMES[i] for i, c in enumerate(bitmask) if c == "1"]

    missing = [_DAY_NAMES[i] for i, c in enumerate(bitmask) if c == "0"]

    count   = len(active)

    if count == 7 or count == 0:

        return "(DAILY)"

    if count in (4, 5):

        # Show the days it does NOT run — shorter and clearer

        return f"{count:02d} DAYS (Except \u2013 {', '.join(missing)})"

    # 1, 2, 3, or 6 days — show the days it runs

    label = "DAY" if count == 1 else "DAYS"

    return f"{count:02d} {label} ({', '.join(active)})"# ==============================================================================

# SECTION 4 - SMART VIA ALGORITHM

# ==============================================================================

def get_smart_via(station_codes: List[str]) -> str:

    if len(station_codes) < 3:

        return "N/A"

    middle = station_codes[1:-1]

    if len(middle) <= 7:

        return ", ".join(middle)

    anchor_first = middle[0]

    anchor_last  = middle[-1]

    inner        = middle[1:-1]

    if len(inner) <= 5:

        chosen = inner[:]

    else:

        step   = (len(inner) - 1) / 4.0

        chosen = [inner[round(i * step)] for i in range(5)]

    return ", ".join([anchor_first] + chosen + [anchor_last])

# ==============================================================================

# SECTION 5 - WORD DOCUMENT UTILITIES

# ==============================================================================

_FONT_NAME = "Times New Roman"

_FONT_SIZE = 13

_HEADER_FONT_SIZE = 14

def _set_cell_valign(cell, align: str = "center") -> None:

    tc   = cell._tc

    tcPr = tc.get_or_add_tcPr()

    vAlign = OxmlElement("w:vAlign")

    vAlign.set(qn("w:val"), align)

    tcPr.append(vAlign)

def _merge_row(table, row_idx: int, from_col: int, to_col: int):

    a = table.rows[row_idx].cells[from_col]

    b = table.rows[row_idx].cells[to_col]

    a.merge(b)

    return a

def _write(cell, text: str, *,

           bold: bool = False,

           align=WD_ALIGN_PARAGRAPH.LEFT,

           new_para: bool = False,

           font_size: Optional[float] = None) -> None:

    if not new_para:

        para = cell.paragraphs[0]

        para.clear()

    else:

        para = cell.add_paragraph()

    para.alignment = align

    run            = para.add_run(text)

    run.bold       = bold

    run.font.name  = _FONT_NAME

    run.font.size  = Pt(font_size if font_size else _FONT_SIZE)

def _write_bold_prefix(cell, prefix: str, val: str, *,

                      align=WD_ALIGN_PARAGRAPH.LEFT,

                      new_para: bool = False,

                      font_size: Optional[float] = None) -> None:

    if not new_para:

        para = cell.paragraphs[0]

        para.clear()

    else:

        para = cell.add_paragraph()

    para.alignment = align

    size = font_size if font_size else _FONT_SIZE

    r1 = para.add_run(prefix)

    r1.bold = True

    r1.font.name = _FONT_NAME

    r1.font.size = Pt(size)

    r2 = para.add_run(val)

    r2.bold = False

    r2.font.name = _FONT_NAME

    r2.font.size = Pt(size)

# ==============================================================================

# SECTION 6 - TABLE BUILDER (Dynamic rows based on type)

# ==============================================================================

def build_pair_table(doc: Document, up: dict, dn: dict, schedule_type: str = "normal", up_upto: str = "", dn_upto: str = "", up_sections: str = "", dn_sections: str = "") -> None:

    """

    Dynamic black-and-white schedule table.

    """

    up_no = up["train_number"]

    dn_no = dn["train_number"]

    raw_name = _clean_train_name(up.get("train_name", ""))

    name_str = f", {raw_name}" if raw_name else ""

    up_orig = up.get('origin_code', '---')

    dn_orig = dn.get('origin_code', '---')

    # Safeguard against identical station codes (e.g. MAS-MAS)

    if up_orig == dn_orig:

        if up.get('dest_code') and up['dest_code'] != up_orig:

            dn_orig = up['dest_code']

        elif dn.get('dest_code') and dn['dest_code'] != up_orig:

            dn_orig = dn['dest_code']

    train_pair_label = (

        f"{up_no}-{dn_no[-2:]}, "

        f"{up_orig}-{dn_orig}{name_str}"

    )

    up_days = up['run_days']

    dn_days = dn['run_days']

    if schedule_type in ("tod", "tod_wcb"):

        up_days = up_days.replace("DAY", "DAYS")

        dn_days = dn_days.replace("DAY", "DAYS")

    freq_up = f"{up_no}- Ex- {up['origin_code']} - {up_days}"

    if schedule_type in ("tod", "tod_wcb") and up_upto and up_upto.strip():

        freq_up += f" UPTO {up_upto.strip()}"

    freq_dn = f"{dn_no}- Ex- {dn['origin_code']} - {dn_days}"

    if schedule_type in ("tod", "tod_wcb") and dn_upto and dn_upto.strip():

        freq_dn += f" UPTO {dn_upto.strip()}"

    via_str = get_smart_via(up["station_codes"])

    rows_count = 5

    if schedule_type in ("sections", "wcb", "tod_wcb"):

        rows_count = 6

    table           = doc.add_table(rows=rows_count, cols=3)

    table.style     = "Table Grid"

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = [Cm(4.0), Cm(6.5), Cm(6.5)]

    for row in table.rows:

        for i, cell in enumerate(row.cells):

            cell.width = col_w[i]

            _set_cell_valign(cell, "center")

    _write(_merge_row(table, 0, 0, 2), "Detail",

           bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _write(table.cell(1, 0), "Train No.& Name")

    _write(_merge_row(table, 1, 1, 2), train_pair_label, bold=True)

    

    current_row = 2

    if schedule_type == "sections":

        _write(table.cell(current_row, 0), "Catering Services to be excluded in following sections")

        sections_cell = _merge_row(table, current_row, 1, 2)

        up_sec_text = f"{up_no}- {up_sections.strip()}" if up_sections and up_sections.strip() else f"{up_no}- "

        dn_sec_text = f"{dn_no}- {dn_sections.strip()}" if dn_sections and dn_sections.strip() else f"{dn_no}- "

        _write(sections_cell, up_sec_text, bold=True)

        _write(sections_cell, dn_sec_text, bold=True, new_para=True)

        current_row += 1

    _write(table.cell(current_row, 0), "Frequency")

    fr_cell = _merge_row(table, current_row, 1, 2)

    _write(fr_cell, freq_up, bold=True)

    _write(fr_cell, freq_dn, bold=True, new_para=True)

    current_row += 1

    

    _write(table.cell(current_row, 0), "Running Between")

    left_cell = table.cell(current_row, 1)

    _write_bold_prefix(left_cell, "Ex- ", up['origin_code'])

    _write_bold_prefix(left_cell, "Dep:- ", up['dep_time'], new_para=True)

    _write_bold_prefix(left_cell, "Arr:- ", dn['arr_time'], new_para=True)

    right_cell = table.cell(current_row, 2)

    _write_bold_prefix(right_cell, "Ex- ", dn['origin_code'])

    _write_bold_prefix(right_cell, "Dep:- ", dn['dep_time'], new_para=True)

    _write_bold_prefix(right_cell, "Arr:- ", up['arr_time'], new_para=True)

    current_row += 1

    

    _write(table.cell(current_row, 0), "Via")

    _write(_merge_row(table, current_row, 1, 2), via_str)

    current_row += 1

    if schedule_type in ("wcb", "tod_wcb"):

        _write(table.cell(current_row, 0), "Coaches")

        coaches_val = up.get("coaches", "20 Coaches")

        _write(_merge_row(table, current_row, 1, 2), coaches_val, bold=True)

def build_service_table_with_title(doc: Document, up_no: str, dn_no: str) -> None:

    """

    Appends an 8-column Kitchen/Service details table for a train pair at the end of the document.

    """

    dn_suffix = dn_no[-2:] if len(dn_no) >= 2 else dn_no

    heading_text = f"Train no.{up_no}-{dn_suffix}"

    p = doc.add_paragraph()

    p.paragraph_format.space_before = Pt(12)

    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(heading_text)

    run.bold = True

    run.underline = True

    run.font.name = _FONT_NAME

    run.font.size = Pt(14)

    table = doc.add_table(rows=2, cols=8)

    table.style = "Table Grid"

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [

        "Day ", "Service", "Station", "Time",

        "Kitchen licensee", "Contact details",

        "Amount of Spl SD to be paid", "IRCTC zone"

    ]

    col_w = [Cm(1.4), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.5), Cm(2.3), Cm(2.4), Cm(1.8)]

    for row in table.rows:

        for i, cell in enumerate(row.cells):

            cell.width = col_w[i]

            _set_cell_valign(cell, "center")

    for i, h in enumerate(headers):

        cell = table.cell(0, i)

        align = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT

        _write(cell, h, bold=True, align=align)

    for i in range(8):

        cell = table.cell(1, i)

        _write(cell, "Nil", bold=True)

# ==============================================================================

# SECTION 7 - REST ENDPOINTS

# ==============================================================================

class TrainPair(BaseModel):

    up: str

    down: str

    schedule_type: Optional[str] = "normal"

    up_upto: Optional[str] = ""

    dn_upto: Optional[str] = ""

    up_sections: Optional[str] = ""

    dn_sections: Optional[str] = ""

class ScheduleRequest(BaseModel):

    pairs: List[TrainPair]

    auth_password: Optional[str] = None

SECURITY_KEY = "Suhail_Apprentice"

def _remove_file(path: str):

    try:

        os.remove(path)

    except Exception:

        pass

# ==============================================================================

# STATS PERSISTENCE  (stats.json lives next to main.py)

# ==============================================================================

STATS_PATH = os.path.join(BASE_DIR, "stats.json")

_STATS_DEFAULT = {"total_generated": 0, "total_trains": 0, "total_time_saved_minutes": 0}

def _load_stats() -> dict:

    try:

        if os.path.exists(STATS_PATH):

            with open(STATS_PATH, "r") as f:

                return json.load(f)

    except (json.decoder.JSONDecodeError, IOError) as e:

        print(f"[WARN] Failed to load stats (file might be corrupted): {e}. Resetting to defaults.")

    except Exception:

        pass

    

    return dict(_STATS_DEFAULT)

def _save_stats(stats: dict) -> None:

    tmp_name = None

    try:

        # Create temp file in the same directory as stats.json

        fd, tmp_name = tempfile.mkstemp(dir=os.path.dirname(STATS_PATH), text=True)

        with os.fdopen(fd, "w") as f:

            json.dump(stats, f, indent=2)

            f.flush()

            os.fsync(f.fileno())  # Ensure bytes are written to physical disk

            

        # Atomically replace old stats file with the new complete one

        os.replace(tmp_name, STATS_PATH)

    except Exception as e:

        print(f"[WARN] Could not save stats: {e}")

        if tmp_name and os.path.exists(tmp_name):

            try:

                os.remove(tmp_name)

            except Exception:

                pass

def _increment_stats(pairs_count: int) -> None:

    """Called once after a successful document generation."""

    stats = _load_stats()

    stats["total_generated"]          += 1

    stats["total_trains"]             += pairs_count

    stats["total_time_saved_minutes"] += pairs_count * 5   # 5 min saved per pair

    _save_stats(stats)

@app.get("/get-stats")

def get_stats():

    return _load_stats()

_is_generating = False

@app.post("/generate-schedule-stream")

async def generate_schedule_stream(req: ScheduleRequest):

    # Server-side un-bypassable security access check

    if not req.auth_password or req.auth_password.strip() != SECURITY_KEY:

        async def denied_gen():

            yield f"data: {json.dumps({'error': '🔒 Security Access Denied: Invalid or Missing Access Key.', 'type': 'warn'})}\n\n"

        return StreamingResponse(denied_gen(), status_code=401, media_type="text/event-stream")

    global _is_generating

    if _is_generating:

        async def busy_gen():

            yield f"data: {json.dumps({'text': '⚠️ [WARN] A schedule generation is already in progress. Please wait.', 'type': 'warn'})}\n\n"

        return StreamingResponse(busy_gen(), media_type="text/event-stream")

    _is_generating = True

    async def stream_generator():

        global _is_generating

        try:

            yield f"data: {json.dumps({'text': '==================================================', 'type': 'info'})}\n\n"

            yield f"data: {json.dumps({'text': '  IndiaRailInfo Scraper Initialised (No Auth Required)', 'type': 'success'})}\n\n"

            yield f"data: {json.dumps({'text': '==================================================', 'type': 'info'})}\n\n"

            doc     = Document()

            section = doc.sections[0]

            section.top_margin    = Cm(2)

            section.bottom_margin = Cm(2)

            section.left_margin   = Cm(2)

            section.right_margin  = Cm(2)

            generated = 0

            processed_pairs = []

            for pair_idx, pair in enumerate(req.pairs):

                up_no, dn_no = pair.up.strip(), pair.down.strip()

                if not up_no or not dn_no:

                    continue

                header_str = f"[PAIR #{pair_idx + 1}] {up_no} ⇄ {dn_no} [{pair.schedule_type.upper()}]"

                print(f"\n{C.B_BLUE}{C.BOLD}┌──────────────────────────────────────────────────────────────┐{C.RESET}")

                print(f"{C.B_BLUE}{C.BOLD}│ 🚆 PROCESSING {header_str:<46} │{C.RESET}")

                print(f"{C.B_BLUE}{C.BOLD}└──────────────────────────────────────────────────────────────┘{C.RESET}")

                pair_log_text = "\n" + header_str

                yield f"data: {json.dumps({'text': pair_log_text, 'type': 'pair'})}\n\n"

                await asyncio.sleep(0.05)

                # UP Train

                up_msg = f"  🔍 [IRI] Fetching Schedule for Train {up_no}..."

                print(f"  {C.B_CYAN}🔍 [IRI]{C.RESET} Fetching Schedule for Train {C.B_YELLOW}{up_no}{C.RESET}...")

                yield f"data: {json.dumps({'text': up_msg, 'type': 'info'})}\n\n"

                await asyncio.sleep(0.05)

                up_info = fetch_train_indiarailinfo(up_no)

                if not up_info:

                    err_msg = f"  ⚠️ [WARN] Train {up_no} not found on IndiaRailInfo."

                    yield f"data: {json.dumps({'text': err_msg, 'type': 'warn'})}\n\n"

                    yield f"data: {json.dumps({'error': f'Train {up_no} not found on IndiaRailInfo.'})}\n\n"

                    return

                up_ok_msg = f"  ✅ [OK] Successfully fetched Train {up_no} ({up_info['train_name']})"

                yield f"data: {json.dumps({'text': up_ok_msg, 'type': 'success'})}\n\n"

                await asyncio.sleep(0.1)

                # DOWN Train

                dn_msg = f"  🔍 [IRI] Fetching Schedule for Train {dn_no}..."

                print(f"  {C.B_CYAN}🔍 [IRI]{C.RESET} Fetching Schedule for Train {C.B_YELLOW}{dn_no}{C.RESET}...")

                yield f"data: {json.dumps({'text': dn_msg, 'type': 'info'})}\n\n"

                await asyncio.sleep(0.05)

                dn_info = fetch_train_indiarailinfo(dn_no)

                if not dn_info:

                    err_msg = f"  ⚠️ [WARN] Train {dn_no} not found on IndiaRailInfo."

                    yield f"data: {json.dumps({'text': err_msg, 'type': 'warn'})}\n\n"

                    yield f"data: {json.dumps({'error': f'Train {dn_no} not found on IndiaRailInfo.'})}\n\n"

                    return

                dn_ok_msg = f"  ✅ [OK] Successfully fetched Train {dn_no} ({dn_info['train_name']})"

                yield f"data: {json.dumps({'text': dn_ok_msg, 'type': 'success'})}\n\n"

                await asyncio.sleep(0.1)

                if generated > 0:

                    p_gap = doc.add_paragraph()

                    p_gap.paragraph_format.space_before = Pt(6)

                    p_gap.paragraph_format.space_after = Pt(12)

                build_pair_table(doc, up_info, dn_info, pair.schedule_type, up_upto=pair.up_upto or "", dn_upto=pair.dn_upto or "", up_sections=pair.up_sections or "", dn_sections=pair.dn_sections or "")

                processed_pairs.append((up_no, dn_no))

                generated += 1

                built_msg = f"  ✨ [SUCCESS] Table Built: {up_info['origin_code']} ⇄ {dn_info['origin_code']} [{pair.schedule_type.upper()}]"

                print(f"  {C.B_GREEN}✨ [SUCCESS]{C.RESET} Table Built: {C.B_WHITE}{up_info['origin_code']}{C.RESET} ⇄ {C.B_WHITE}{dn_info['origin_code']}{C.RESET} {C.DIM}[{pair.schedule_type.upper()}]{C.RESET}")

                yield f"data: {json.dumps({'text': built_msg, 'type': 'success'})}\n\n"

                await asyncio.sleep(0.1)

            if processed_pairs:

                service_log_text = "\n📋 Appending Kitchen/Service details tables at document end..."

                yield f"data: {json.dumps({'text': service_log_text, 'type': 'info'})}\n\n"

                for up_no, dn_no in processed_pairs:

                    build_service_table_with_title(doc, up_no, dn_no)

            if generated == 0:

                yield f"data: {json.dumps({'error': 'No valid train pairs could be processed.'})}\n\n"

                return

            fname = f"schedule_{uuid.uuid4().hex}.docx"

            out_path = os.path.join(BASE_DIR, fname)

            doc.save(out_path)
            print(f"[RENDER LOG] 📦 Word Document saved: {out_path} ({os.path.getsize(out_path)} bytes)", flush=True)

            _increment_stats(generated)

            comp_text = f"🎉 [SUCCESS] IRCTC Tender Schedule Word Document Generated! ({generated} Pair{'s' if generated > 1 else ''})"

            print(f"\n{C.B_GREEN}{C.BOLD}============================================================={C.RESET}")

            print(f" {C.B_GREEN}{C.BOLD}🎉 DOCUMENT GENERATION COMPLETED SUCCESSFULLY!{C.RESET}")

            print(f" {C.B_WHITE}📦 Output File:{C.RESET} {C.B_YELLOW}{fname}{C.RESET}")

            print(f" {C.B_CYAN}📊 Main Schedules:{C.RESET} {generated}  |  {C.B_MAGENTA}📋 Service Tables:{C.RESET} {len(processed_pairs)} (at End)")

            print(f"{C.B_GREEN}{C.BOLD}============================================================={C.RESET}\n")

            yield f"data: {json.dumps({'text': comp_text, 'type': 'success', 'done': True, 'filename': fname})}\n\n"

        finally:

            _is_generating = False

    return StreamingResponse(stream_generator(), media_type="text/event-stream")

@app.get("/download-file/{filename}")

def download_file(filename: str, background_tasks: BackgroundTasks):
    print(f"[RENDER LOG] 📤 Serving download for {filename}...", flush=True)

    safe_name = os.path.basename(filename)

    file_path = os.path.join(BASE_DIR, safe_name)

    if not os.path.exists(file_path):

        raise HTTPException(status_code=404, detail="Requested schedule file not found.")

    background_tasks.add_task(_remove_file, file_path)

    return FileResponse(

        path=file_path,

        filename="IRCTC_Tender_Schedules.docx",

        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )

@app.post("/extract-trains")

async def extract_trains(file: UploadFile = File(...)):

    """

    PDF parser with robust regex for all train number formats:

      - 20151-52   -> UP: 20151, DOWN: 20152  (2-digit suffix)

      - 13247-13248 -> UP: 13247, DOWN: 13248  (full 5-digit pair)

      - 12601       -> singleton (try to pair with next sequential)

    """

    if not file.filename.lower().endswith(".pdf"):

        raise HTTPException(status_code=400, detail="File must be a PDF.")

    contents = await file.read()

    text = ""

    # ── Extractor 1: PyMuPDF (fitz) — handles non-standard font encodings ────

    # Much better than pdfplumber for government/office-generated PDFs.

    try:

        import fitz  # PyMuPDF

        doc_fitz = fitz.open(stream=contents, filetype="pdf")

        for page in doc_fitz:

            # get_text("text") — plain text with newlines

            page_text = page.get_text("text")

            if page_text.strip():

                text += page_text + "\n"

            # get_text("blocks") — text blocks with exact layout

            # Each block: (x0,y0,x1,y1, "text", block_no, block_type)

            for block in page.get_text("blocks"):

                block_text = block[4] if len(block) > 4 else ""

                if block_text.strip():

                    text += block_text.replace("\n", " ") + " "

            # get_text("dict") — span-level detail, catches chars in odd fonts

            block_dict = page.get_text("dict")

            for blk in block_dict.get("blocks", []):

                for line in blk.get("lines", []):

                    line_str = "".join(span["text"] for span in line.get("spans", []))

                    if line_str.strip():

                        text += line_str + " "

            text += "\n"

        doc_fitz.close()

    except Exception as e:

        print(f"  [WARN] PyMuPDF failed: {e}. Falling back to pdfplumber.")

    # ── Extractor 2: pdfplumber — fallback if PyMuPDF got nothing ───────────

    if not text.strip():

        try:

            with pdfplumber.open(io.BytesIO(contents)) as pdf:

                for page in pdf.pages:

                    page_text = page.extract_text()

                    if page_text:

                        text += page_text + "\n"

                    try:

                        tables = page.extract_tables()

                        for tbl in (tables or []):

                            for row in tbl:

                                for cell in row:

                                    if cell:

                                        text += cell.strip().replace("\n", " ") + " "

                            text += "\n"

                    except Exception:

                        pass

        except Exception as exc:

            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {exc}")

    # ── Extractor 3: OCR via easyocr — for scanned image PDFs ───────────────

    # Triggered only when both text-based extractors return nothing.

    # PyMuPDF renders each PDF page as a 300-DPI image; easyocr reads it.

    if not text.strip():

        print("[OCR] No text layer found. Running OCR on page images...")

        try:

            import fitz

            import numpy as np

            reader = _get_ocr_reader()

            doc_fitz = fitz.open(stream=contents, filetype="pdf")

            for page_num, page in enumerate(doc_fitz):

                # ── Render at 200 DPI (sufficient for large printed numbers) ─

                mat = fitz.Matrix(200 / 72, 200 / 72)

                pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)

                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(

                          pix.height, pix.width, 3)

                h, w, _ = img.shape

                

                # ── Dynamic Cropping based on Orientation ────────────────────

                if h > w:

                    orientation = "Portrait"

                    crop_pct = 0.32

                else:

                    orientation = "Landscape"

                    crop_pct = 0.20

                    

                crop_w = max(1, int(w * crop_pct))

                img_crop = img[:, :crop_w, :]

                

                print(f"[OCR] Page {page_num+1} detected as [{orientation}]. Cropping at {int(crop_pct*100)}%...")

                results = reader.readtext(img_crop, detail=0, paragraph=False)

                page_text = " ".join(results)

                

                # ── Safety Net: Full Page Scan Fallback ──────────────────────

                if not re.search(r'\d{5}', page_text):

                    print(f"  [WARN] No train numbers found in {int(crop_pct*100)}% crop. Triggering Full Page Scan...")

                    results = reader.readtext(img, detail=0, paragraph=False)

                    page_text = " ".join(results)

                print(f"  [OCR] Extracted preview: {repr(page_text[:200])}")

                text += page_text + "\n"

            doc_fitz.close()

        except Exception as ocr_err:

            raise HTTPException(

                status_code=500,

                detail=f"PDF has no text layer and OCR also failed: {ocr_err}"

            )

    # Normalise dash variants and invisible chars

    text = (text

            .replace("\u2013", "-")   # en-dash

            .replace("\u2014", "-")   # em-dash

            .replace("\u2011", "-")   # non-breaking hyphen

            .replace("\u00a0", " ")   # non-breaking space

            .replace("\ufb01", "fi")  # ligature fi

            )

    print(f"[PDF] Chars extracted: {len(text)}")

    print(f"[PDF] Text sample: {repr(text[:400])}")

    pairs:      List[List[str]] = []

    used_pairs: set             = set()

    singletons: List[str]       = []

    # --- Pass 1: Full 5-digit pairs  e.g. "13247-13248" or "13247/13248" -----

    full_pair_re = re.compile(r"\b(\d{5})[-/](\d{5})\b")

    for m in full_pair_re.finditer(text):

        a, b = m.group(1), m.group(2)

        key = (a, b)

        if key not in used_pairs:

            pairs.append([a, b])

            used_pairs.add(key)

    # Remove matched positions so Pass 2 doesn't double-count

    text_reduced = full_pair_re.sub("     ", text)   # blank out with spaces

    # --- Pass 2: Suffix pairs  e.g. "20151-52" or "16779-80" -----------------

    # Suffix can be 2-4 digits; DOWN = UP[:(5-len(suffix))] + suffix

    suffix_pair_re = re.compile(r"\b(\d{5})[-/](\d{2,4})\b")

    for m in suffix_pair_re.finditer(text_reduced):

        base   = m.group(1)

        suffix = m.group(2)

        # Build full 5-digit DOWN train number

        prefix_len = 5 - len(suffix)

        down = base[:prefix_len] + suffix

        # Sanity check: result must be 5 digits

        if len(down) != 5 or not down.isdigit():

            continue

        key = (base, down)

        if key not in used_pairs:

            pairs.append([base, down])

            used_pairs.add(key)

    text_reduced = suffix_pair_re.sub("     ", text_reduced)

    # --- Pass 3: Remaining standalone 5-digit numbers -------------------------

    single_re = re.compile(r"\b(\d{5})\b")

    for m in single_re.finditer(text_reduced):

        tn = m.group(1)

        # Skip numbers already captured as part of a pair

        if any(tn == p[0] or (len(p) > 1 and tn == p[1]) for p in pairs):

            continue

        singletons.append(tn)

    # Try to group singletons into sequential pairs (e.g. 12601 + 12602)

    unique_singles = sorted(set(singletons))

    used_singles:  set = set()

    for tn in unique_singles:

        if tn in used_singles:

            continue

        nxt = str(int(tn) + 1).zfill(5)

        if nxt in unique_singles and nxt not in used_singles:

            key = (tn, nxt)

            if key not in used_pairs:

                pairs.append([tn, nxt])

                used_pairs.add(key)

            used_singles.update([tn, nxt])

        else:

            # True singleton — include anyway (user can remove from UI)

            pairs.append([tn])

            used_singles.add(tn)

    print(f"[PDF] Extracted {len(pairs)} pairs from '{file.filename}'")

    for p in pairs:

        print(f"       {p}")

    return {"extracted_pairs": pairs}

# ==============================================================================

# FRONTEND STATIC FILES MOUNT (Single-Server Deployment)

# ==============================================================================

FRONTEND_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "frontend"))

if os.path.exists(FRONTEND_DIR):

    app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")

if __name__ == "__main__":

    import uvicorn

    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)

